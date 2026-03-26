import cv2
import numpy as np
import pytesseract
import easyocr
from PIL import Image, ImageEnhance, ImageFilter, ImageOps
import fitz  # PyMuPDF
import docx
import tempfile
import json
import logging
from typing import Tuple, Dict, Any, Optional, List
import re
from dataclasses import dataclass
from pathlib import Path
import io
import base64

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class OCRResult:
    """Data class for OCR results"""
    
    text: str
    confidence: float
    method: str
    preprocessing_applied: List[str]
    extraction_time: float

class AdvancedImagePreprocessor:
    """Advanced image preprocessing for optimal OCR results"""
    
    @staticmethod
    def deskew_image(image: np.ndarray) -> np.ndarray:
        """Detect and correct skew in the image"""
        try:
            # Convert to grayscale if needed
            if len(image.shape) == 3:
                gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            else:
                gray = image.copy()
            
            # Find all non-zero pixels (text pixels)
            coords = np.column_stack(np.where(gray > 0))
            
            if coords.shape[0] < 10:  # Not enough points for reliable skew detection
                return image
            
            # Get the minimum area rectangle
            angle = cv2.minAreaRect(coords)[-1]
            
            # Correct the angle
            if angle < -45:
                angle = -(90 + angle)
            else:
                angle = -angle
            
            # Only apply rotation if angle is significant (> 0.5 degrees)
            if abs(angle) > 0.5:
                (h, w) = image.shape[:2]
                center = (w // 2, h // 2)
                M = cv2.getRotationMatrix2D(center, angle, 1.0)
                rotated = cv2.warpAffine(image, M, (w, h), 
                                       flags=cv2.INTER_CUBIC, 
                                       borderMode=cv2.BORDER_REPLICATE)
                return rotated
            
            return image
            
        except Exception as e:
            logger.warning(f"Deskewing failed: {e}")
            return image
    
    @staticmethod
    def remove_noise(image: np.ndarray) -> np.ndarray:
        """Remove noise from the image"""
        try:
            # Apply Non-local Means Denoising
            if len(image.shape) == 3:
                denoised = cv2.fastNlMeansDenoisingColored(image, None, 10, 10, 7, 21)
            else:
                denoised = cv2.fastNlMeansDenoising(image, None, 10, 7, 21)
            
            # Morphological operations to remove small noise
            kernel = np.ones((2, 2), np.uint8)
            denoised = cv2.morphologyEx(denoised, cv2.MORPH_OPEN, kernel)
            
            return denoised
        except Exception as e:
            logger.warning(f"Noise removal failed: {e}")
            return image
    
    @staticmethod
    def enhance_contrast_adaptive(image: np.ndarray) -> np.ndarray:
        """Apply adaptive histogram equalization for better contrast"""
        try:
            if len(image.shape) == 3:
                # Convert to LAB color space
                lab = cv2.cvtColor(image, cv2.COLOR_BGR2LAB)
                l, a, b = cv2.split(lab)
                
                # Apply CLAHE to L channel
                clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
                l = clahe.apply(l)
                
                # Merge channels and convert back
                enhanced = cv2.merge([l, a, b])
                enhanced = cv2.cvtColor(enhanced, cv2.COLOR_LAB2BGR)
            else:
                # Apply CLAHE directly for grayscale
                clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
                enhanced = clahe.apply(image)
            
            return enhanced
        except Exception as e:
            logger.warning(f"Contrast enhancement failed: {e}")
            return image
    
    @staticmethod
    def apply_threshold(image: np.ndarray, method: str = "adaptive") -> np.ndarray:
        """Apply various thresholding techniques"""
        try:
            if len(image.shape) == 3:
                gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            else:
                gray = image.copy()
            
            if method == "adaptive":
                # Adaptive Gaussian thresholding
                thresh = cv2.adaptiveThreshold(
                    gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                    cv2.THRESH_BINARY, 11, 2
                )
            elif method == "otsu":
                # Otsu's thresholding
                _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            elif method == "mean":
                # Adaptive mean thresholding
                thresh = cv2.adaptiveThreshold(
                    gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C, 
                    cv2.THRESH_BINARY, 11, 2
                )
            else:
                thresh = gray
            
            return thresh
        except Exception as e:
            logger.warning(f"Thresholding failed: {e}")
            return image
    
    @staticmethod
    def resize_for_ocr(image: np.ndarray, target_height: int = 2000) -> np.ndarray:
        """Resize image to optimal size for OCR"""
        try:
            height, width = image.shape[:2]
            
            # Only resize if image is too small
            if height < target_height:
                scale_factor = target_height / height
                new_width = int(width * scale_factor)
                new_height = target_height
                
                resized = cv2.resize(image, (new_width, new_height), 
                                   interpolation=cv2.INTER_CUBIC)
                return resized
            
            return image
        except Exception as e:
            logger.warning(f"Resizing failed: {e}")
            return image

class EnhancedOCREngine:
    """Enhanced OCR engine with multiple extraction methods"""
    
    def __init__(self):
        self.preprocessor = AdvancedImagePreprocessor()
        # Initialize EasyOCR reader
        try:
            self.easyocr_reader = easyocr.Reader(['en'], gpu=False)
        except Exception as e:
            logger.warning(f"EasyOCR initialization failed: {e}")
            self.easyocr_reader = None
    
    def preprocess_image(self, pil_image: Image.Image, aggressive: bool = False) -> Tuple[np.ndarray, List[str]]:
        """Comprehensive image preprocessing pipeline"""
        applied_methods = []
        
        # Convert PIL to OpenCV format
        if pil_image.mode == 'RGBA':
            pil_image = pil_image.convert('RGB')
        
        img_array = np.array(pil_image)
        if len(img_array.shape) == 3:
            img_array = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
        
        original_img = img_array.copy()
        
        try:
            # Step 1: Resize for better OCR
            img_array = self.preprocessor.resize_for_ocr(img_array)
            applied_methods.append("resize")
            
            # Step 2: Deskew the image
            img_array = self.preprocessor.deskew_image(img_array)
            applied_methods.append("deskew")
            
            # Step 3: Enhance contrast
            img_array = self.preprocessor.enhance_contrast_adaptive(img_array)
            applied_methods.append("contrast_enhancement")
            
            # Step 4: Remove noise
            if aggressive:
                img_array = self.preprocessor.remove_noise(img_array)
                applied_methods.append("noise_removal")
            
            # Step 5: Apply thresholding
            img_array = self.preprocessor.apply_threshold(img_array, "adaptive")
            applied_methods.append("adaptive_threshold")
            
        except Exception as e:
            logger.error(f"Preprocessing failed: {e}")
            return original_img, ["error_fallback"]
        
        return img_array, applied_methods
    
    def extract_with_tesseract(self, image: np.ndarray, config_variations: bool = True) -> OCRResult:
        """Extract text using Tesseract with multiple configurations"""
        import time
        start_time = time.time()
        
        try:
            # Convert OpenCV image to PIL for Tesseract
            if len(image.shape) == 3:
                pil_img = Image.fromarray(cv2.cvtColor(image, cv2.COLOR_BGR2RGB))
            else:
                pil_img = Image.fromarray(image)
            
            best_text = ""
            best_confidence = 0
            best_config = ""
            
            # Different PSM configurations to try
            psm_configs = [
                (6, '--oem 3 --psm 6'),  # Uniform block of text
                (3, '--oem 3 --psm 3'),  # Fully automatic page segmentation
                (4, '--oem 3 --psm 4'),  # Single column of text
                (7, '--oem 3 --psm 7'),  # Single text line
                (8, '--oem 3 --psm 8'),  # Single word
                (11, '--oem 3 --psm 11'), # Sparse text
                (13, '--oem 3 --psm 13')  # Raw line
            ]
            
            for psm, config in psm_configs:
                try:
                    # Extract text
                    text = pytesseract.image_to_string(pil_img, config=config).strip()
                    
                    # Get confidence data
                    data = pytesseract.image_to_data(pil_img, config=config, output_type=pytesseract.Output.DICT)
                    confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
                    avg_confidence = np.mean(confidences) if confidences else 0
                    
                    # Select best result based on text length and confidence
                    score = len(text) * 0.7 + avg_confidence * 0.3
                    current_score = len(best_text) * 0.7 + best_confidence * 0.3
                    
                    if score > current_score:
                        best_text = text
                        best_confidence = avg_confidence
                        best_config = config
                    
                    # If we have good results, don't try all configurations
                    if not config_variations and len(text) > 50 and avg_confidence > 70:
                        break
                        
                except Exception as e:
                    logger.debug(f"Tesseract PSM {psm} failed: {e}")
                    continue
            
            extraction_time = time.time() - start_time
            
            return OCRResult(
                text=best_text,
                confidence=best_confidence,
                method=f"tesseract_{best_config}",
                preprocessing_applied=[],
                extraction_time=extraction_time
            )
            
        except Exception as e:
            logger.error(f"Tesseract extraction failed: {e}")
            return OCRResult("", 0, "tesseract_error", [], time.time() - start_time)
    
    def extract_with_easyocr(self, image: np.ndarray) -> OCRResult:
        """Extract text using EasyOCR"""
        import time
        start_time = time.time()
        
        if self.easyocr_reader is None:
            return OCRResult("", 0, "easyocr_unavailable", [], 0)
        
        try:
            # EasyOCR expects RGB format
            if len(image.shape) == 3:
                rgb_image = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
            else:
                rgb_image = cv2.cvtColor(image, cv2.COLOR_GRAY2RGB)
            
            results = self.easyocr_reader.readtext(rgb_image, detail=1)
            
            # Extract text and calculate average confidence
            texts = []
            confidences = []
            
            for (bbox, text, confidence) in results:
                if confidence > 0.3:  # Filter low confidence results
                    texts.append(text)
                    confidences.append(confidence * 100)
            
            combined_text = ' '.join(texts)
            avg_confidence = np.mean(confidences) if confidences else 0
            
            extraction_time = time.time() - start_time
            
            return OCRResult(
                text=combined_text,
                confidence=avg_confidence,
                method="easyocr",
                preprocessing_applied=[],
                extraction_time=extraction_time
            )
            
        except Exception as e:
            logger.error(f"EasyOCR extraction failed: {e}")
            return OCRResult("", 0, "easyocr_error", [], time.time() - start_time)
    
    def extract_text_comprehensive(self, pil_image: Image.Image) -> OCRResult:
        """Comprehensive text extraction using multiple methods"""
        best_result = OCRResult("", 0, "none", [], 0)
        
        # Try different preprocessing approaches
        preprocessing_approaches = [
            (False, "standard"),
            (True, "aggressive")
        ]
        
        for aggressive, approach_name in preprocessing_approaches:
            try:
                # Preprocess image
                processed_img, applied_methods = self.preprocess_image(pil_image, aggressive=aggressive)
                
                # Try Tesseract
                tesseract_result = self.extract_with_tesseract(processed_img)
                tesseract_result.preprocessing_applied = applied_methods
                tesseract_result.method += f"_{approach_name}"
                
                # Try EasyOCR
                easyocr_result = self.extract_with_easyocr(processed_img)
                easyocr_result.preprocessing_applied = applied_methods
                easyocr_result.method += f"_{approach_name}"
                
                # Select best result
                candidates = [tesseract_result, easyocr_result]
                
                for result in candidates:
                    # Score based on text length and confidence
                    score = len(result.text) * 0.6 + result.confidence * 0.4
                    best_score = len(best_result.text) * 0.6 + best_result.confidence * 0.4
                    
                    if score > best_score:
                        best_result = result
                
                # If we have good results, don't try more aggressive preprocessing
                if len(best_result.text) > 100 and best_result.confidence > 75:
                    break
                    
            except Exception as e:
                logger.error(f"Comprehensive extraction failed for {approach_name}: {e}")
                continue
        
        return best_result

class EnhancedInvoiceProcessor:
    """Enhanced invoice processing with multiple file format support"""
    
    def __init__(self):
        self.ocr_engine = EnhancedOCREngine()
        
    def extract_from_image(self, uploaded_file) -> Tuple[str, str, Dict]:
        """Extract text from image files with enhanced OCR"""
        try:
            image = Image.open(uploaded_file)
            
            # Convert RGBA to RGB if necessary
            if image.mode == 'RGBA':
                # Create white background
                background = Image.new('RGB', image.size, (255, 255, 255))
                background.paste(image, mask=image.split()[-1])
                image = background
            
            # Perform comprehensive OCR
            ocr_result = self.ocr_engine.extract_text_comprehensive(image)
            
            # Prepare metadata
            metadata = {
                'ocr_method': ocr_result.method,
                'confidence': ocr_result.confidence,
                'preprocessing': ocr_result.preprocessing_applied,
                'extraction_time': ocr_result.extraction_time,
                'text_length': len(ocr_result.text)
            }
            
            return ocr_result.text, 'image', metadata
            
        except Exception as e:
            logger.error(f"Image extraction failed: {e}")
            return f"[Image Extraction Error: {e}]", 'image_error', {}
    
    def extract_from_pdf(self, uploaded_file) -> Tuple[str, str, Dict]:
        """Extract text from PDF files"""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(uploaded_file.read())
                tmp_file.flush()
                
                doc = fitz.open(tmp_file.name)
                extracted_texts = []
                
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    
                    # Try text extraction first
                    text = page.get_text()
                    
                    # If no text found, try OCR on page image
                    if len(text.strip()) < 50:
                        try:
                            # Render page as image
                            mat = fitz.Matrix(2, 2)  # 2x zoom
                            pix = page.get_pixmap(matrix=mat)
                            img_data = pix.tobytes("ppm")
                            
                            # Convert to PIL Image
                            pil_image = Image.open(io.BytesIO(img_data))
                            
                            # Apply OCR
                            ocr_result = self.ocr_engine.extract_text_comprehensive(pil_image)
                            text = ocr_result.text
                            
                        except Exception as ocr_error:
                            logger.warning(f"OCR on PDF page {page_num} failed: {ocr_error}")
                    
                    if text.strip():
                        extracted_texts.append(text)
                
                doc.close()
                
                combined_text = "\n\n".join(extracted_texts)
                metadata = {
                    'pages_processed': len(doc),
                    'text_length': len(combined_text),
                    'extraction_method': 'pdf_hybrid'
                }
                
                return combined_text, 'pdf', metadata
                
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return f"[PDF Extraction Error: {e}]", 'pdf_error', {}
    
    def extract_from_docx(self, uploaded_file) -> Tuple[str, str, Dict]:
        """Extract text from DOCX files"""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(uploaded_file.read())
                tmp_file.flush()
                
                doc = docx.Document(tmp_file.name)
                
                # Extract text from paragraphs
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                
                # Extract text from tables
                table_texts = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            table_texts.append(" | ".join(row_text))
                
                # Combine all text
                all_text = paragraphs + table_texts
                combined_text = "\n".join(all_text)
                
                metadata = {
                    'paragraphs_count': len(paragraphs),
                    'tables_count': len(doc.tables),
                    'text_length': len(combined_text),
                    'extraction_method': 'docx_native'
                }
                
                return combined_text, 'docx', metadata
                
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return f"[DOCX Extraction Error: {e}]", 'docx_error', {}
    
    def extract_from_txt(self, uploaded_file) -> Tuple[str, str, Dict]:
        """Extract text from plain text files"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    uploaded_file.seek(0)  # Reset file pointer
                    text = uploaded_file.read().decode(encoding)
                    
                    metadata = {
                        'encoding_used': encoding,
                        'text_length': len(text),
                        'extraction_method': 'txt_native'
                    }
                    
                    return text, 'txt', metadata
                    
                except UnicodeDecodeError:
                    continue
            
            # Fallback: read with error handling
            uploaded_file.seek(0)
            text = uploaded_file.read().decode('utf-8', errors='ignore')
            
            metadata = {
                'encoding_used': 'utf-8_with_errors_ignored',
                'text_length': len(text),
                'extraction_method': 'txt_fallback'
            }
            
            return text, 'txt', metadata
            
        except Exception as e:
            logger.error(f"TXT extraction failed: {e}")
            return f"[TXT Extraction Error: {e}]", 'txt_error', {}
    
    def extract_invoice_text(self, uploaded_file) -> Tuple[str, str, Dict]:
        """Main method to extract text from various file formats"""
        try:
            file_type = uploaded_file.type
            logger.info(f"Processing file of type: {file_type}")
            
            if file_type.startswith('image/'):
                return self.extract_from_image(uploaded_file)
            elif file_type == 'application/pdf':
                return self.extract_from_pdf(uploaded_file)
            elif file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                              'application/msword']:
                return self.extract_from_docx(uploaded_file)
            elif file_type == 'text/plain':
                return self.extract_from_txt(uploaded_file)
            else:
                logger.warning(f"Unsupported file type: {file_type}")
                return f"[Unsupported file type: {file_type}]", 'unsupported', {}
                
        except Exception as e:
            logger.error(f"General extraction error: {e}")
            return f"[General Extraction Error: {e}]", 'general_error', {}

# Enhanced invoice details extraction with better error handling
def extract_invoice_details_with_llm(invoice_text: str, llm_chain=None) -> Dict[str, Any]:
    """Enhanced invoice extraction with better structure and error handling"""
    try:
        if not llm_chain:
            # Return a more detailed fallback structure
            return {
                "vendor_name": "Not found",
                "invoice_number": "Not found", 
                "date": "Not found",
                "state": "Not found",
                "city": "Not found",
                "amount": "Not found",
                "payment_type": "Not found",
                "confidence": 0,
                "extraction_method": "fallback",
                "text_preview": invoice_text[:200] + "..." if len(invoice_text) > 200 else invoice_text
            }
        
        # Enhanced prompt for better extraction
        enhanced_prompt = f"""
        Analyze the following invoice/document text and extract key information.
        
        Text to analyze:
        {invoice_text[:4000]}
        
        Extract the following information and respond in JSON format:
        {{
            "vendor_name": "Company or vendor name",
            "invoice_number": "Invoice/bill number",
            "date": "Invoice date (YYYY-MM-DD format if possible)",
            "state": "State/province",
            "city": "City name",
            "amount": "Total amount with currency",
            "payment_type": "Description of goods/services",
            "confidence": "Confidence score 0-100",
            "additional_details": "Any other relevant information"
        }}
        
        Use "Not found" for missing information. Be precise and factual.
        """
        
        response = llm_chain.invoke({"prompt": enhanced_prompt})
        
        # Enhanced JSON parsing with fallback
        try:
            cleaned_response = response.replace('```json', '').replace('```', '').strip()
            
            # Try to find JSON in the response
            json_match = re.search(r'\{.*\}', cleaned_response, re.DOTALL)
            if json_match:
                cleaned_response = json_match.group(0)
            
            invoice_data = json.loads(cleaned_response)
            
            # Ensure all required fields exist
            required_fields = ["vendor_name", "invoice_number", "date", "state", "city", "amount", "payment_type", "confidence"]
            for field in required_fields:
                if field not in invoice_data:
                    invoice_data[field] = "Not found"
            
            # Add metadata
            invoice_data["extraction_method"] = "llm_enhanced"
            invoice_data["text_length"] = len(invoice_text)
            
            return invoice_data
            
        except json.JSONDecodeError as json_error:
            logger.error(f"JSON parsing failed: {json_error}")
            # Return structured fallback with partial information
            return {
                "vendor_name": "Extraction failed",
                "invoice_number": "Extraction failed",
                "date": "Extraction failed", 
                "state": "Extraction failed",
                "city": "Extraction failed",
                "amount": "Extraction failed",
                "payment_type": "Extraction failed",
                "confidence": 10,
                "extraction_method": "llm_partial_failure",
                "raw_response": response[:500]
            }
            
    except Exception as e:
        logger.error(f"LLM extraction failed: {e}")
        return {
            "vendor_name": "Error in extraction",
            "invoice_number": "Error in extraction",
            "date": "Error in extraction",
            "state": "Error in extraction", 
            "city": "Error in extraction",
            "amount": "Error in extraction",
            "payment_type": "Error in extraction",
            "confidence": 0,
            "extraction_method": "error",
            "error_message": str(e)
        }

def format_invoice_info(invoice_data: Dict[str, Any], include_metadata: bool = True) -> str:
    """Enhanced formatting of invoice information"""
    if not invoice_data:
        return "No invoice information available."
    
    # Core information
    info_sections = [
        "**📄 Invoice/Document Information:**",
        f"• **Vendor:** {invoice_data.get('vendor_name', 'Not available')}",
        f"• **Invoice Number:** {invoice_data.get('invoice_number', 'Not available')}",
        f"• **Date:** {invoice_data.get('date', 'Not available')}",
        f"• **Amount:** {invoice_data.get('amount', 'Not available')}",
        "",
        "**📍 Location Information:**",
        f"• **State:** {invoice_data.get('state', 'Not available')}",
        f"• **City:** {invoice_data.get('city', 'Not available')}",
        "",
        "**💼 Service Details:**",
        f"• **Payment Type/Services:** {invoice_data.get('payment_type', 'Not available')}",
    ]
    
    # Add metadata if requested
    if include_metadata:
        info_sections.extend([
            "",
            "**🔍 Extraction Details:**",
            f"• **Confidence:** {invoice_data.get('confidence', 'Not available')}%",
            f"• **Method:** {invoice_data.get('extraction_method', 'Not available')}",
        ])
        
        # Add additional details if available
        if 'additional_details' in invoice_data and invoice_data['additional_details'] != 'Not found':
            info_sections.extend([
                f"• **Additional Info:** {invoice_data['additional_details']}"
            ])
    
    return "\n".join(info_sections)

# Example usage and testing functions
def test_ocr_system():
    """Test function for the OCR system"""
    processor = EnhancedInvoiceProcessor()
    
    # Test with different file types
    print("Enhanced OCR Invoice Processing System Ready!")
    print("Supported formats: Images (PNG, JPG, etc.), PDF, DOCX, TXT")
    print("Features: Advanced preprocessing, Multiple OCR engines, Comprehensive error handling")

if __name__ == "__main__":
    test_ocr_system()